"""
Generates full project documentation as a DOCX file.
Run with: python _generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

DARK_BLUE  = RGBColor(0x1A, 0x3C, 0x6E)
MID_BLUE   = RGBColor(0x23, 0x6F, 0xBF)
TEXT_DARK  = RGBColor(0x1F, 0x23, 0x28)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HDR  = RGBColor(0x1A, 0x3C, 0x6E)
TABLE_ROW  = RGBColor(0xF5, 0xF8, 0xFF)


def shade_cell(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def heading(text, level=1):
    p   = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.bold = True
    run.font.size = Pt(20 - (level - 1) * 3)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p


def body(text, bold=False, italic=False, size=11):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text, level=0):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_DARK
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def code_block(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x40, 0x60)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF2F7')
    pPr.append(shd)
    return p


def add_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], TABLE_HDR)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        row_cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            if r_idx % 2 == 0:
                shade_cell(row_cells[c_idx], TABLE_ROW)
            run = row_cells[c_idx].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def hr():
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), '23558F')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(80)
r = cover.add_run("AWS DevOps RAG Agent")
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Full Project Documentation — IBM WatsonX Edition v2")
r2.font.size = Pt(16)
r2.font.color.rgb = MID_BLUE

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)
r3.font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", level=1)
toc_entries = [
    ("1.",    "Project Overview"),
    ("2.",    "Architecture"),
    ("3.",    "Project Structure"),
    ("4.",    "Installation & Quick Start"),
    ("5.",    "Configuration Reference"),
    ("6.",    "Core Modules"),
    ("  6.1", "Config (config.py)"),
    ("  6.2", "Data Models (models.py)"),
    ("  6.3", "Monitoring Collector (monitoring/collector.py)"),
    ("  6.4", "Analyzer Engine (analyzer/engine.py)"),
    ("  6.5", "Self-Healing Executor (healing/executor.py)"),
    ("  6.6", "RAG Vector Store (rag/vector_store.py)"),
    ("  6.7", "Agent Orchestrator (agent/orchestrator.py)"),
    ("  6.8", "Slack Notifications (notifications/slack.py)"),
    ("  6.9", "FastAPI Server (server.py)"),
    ("  6.10","Entry Point (main.py)"),
    ("7.",    "Anomaly Detection Modes"),
    ("8.",    "Monitored Resources & Metrics"),
    ("9.",    "Self-Healing Actions (21 total)"),
    ("10.",   "API Reference"),
    ("11.",   "WatsonX Orchestrate Integration"),
    ("12.",   "Feedback & Learning Loop"),
    ("13.",   "Knowledge Base & Runbooks"),
    ("14.",   "Slack Notification Setup"),
    ("15.",   "Safety & Operational Guardrails"),
    ("16.",   "Extending the Agent"),
    ("17.",   "Dependencies"),
    ("18.",   "Data Models Reference"),
]
for num, title in toc_entries:
    p   = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(11)
    run.font.color.rgb = MID_BLUE if num.strip().endswith('.') else TEXT_DARK
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview")
hr()
body(
    "AWS DevOps RAG Agent is an autonomous, AI-powered Site Reliability Engineering (SRE) "
    "agent that continuously monitors AWS infrastructure, detects anomalies with three "
    "selectable detection modes (threshold, LLM-only, hybrid), reasons about root causes "
    "using a Large Language Model augmented with RAG, and automatically executes 21 self-healing "
    "actions against AWS APIs — all without human intervention. The agent learns from operator "
    "feedback and integrates with WatsonX Orchestrate as conversational skills."
)
doc.add_paragraph()
body("Key capabilities:", bold=True)
bullet("Continuous polling of EC2, RDS, ECS, Lambda, and ALB resources via CloudWatch")
bullet("Three anomaly detection modes: threshold rules, LLM-only (no fixed thresholds), and hybrid")
bullet("LLM-driven root-cause analysis with RAG-retrieved runbook context")
bullet("21 self-healing actions across EC2, RDS, ECS, ALB, and Lambda")
bullet("Feedback learning loop: operator fixes are appended to knowledge base and FAISS rebuilt")
bullet("Slack webhook notifications on anomaly detection and healing completion")
bullet("WatsonX Orchestrate integration: 3 skills importable via /openapi-orchestrate.json")
bullet("DRY_RUN mode for safe audit without executing real actions")
bullet("Real-time WebSocket event streaming to built-in HTML dashboard")
bullet("LLM provider support: OpenAI, AWS Bedrock, Groq, Ollama, IBM WatsonX")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Architecture")
hr()
body(
    "The agent is built on a LangGraph state machine with four nodes: "
    "MONITOR -> ANALYZE -> HEAL -> REPORT, running in an infinite async loop. "
    "Detection mode, LLM provider, and embedding provider are all runtime-configurable."
)
doc.add_paragraph()
body("LangGraph State Machine Nodes:", bold=True)
bullet("MONITOR  — Collects ResourceSnapshot objects from all 5 AWS services via CloudWatch + APIs")
bullet("ANALYZE  — Detects anomalies (threshold / LLM / hybrid), sends Slack alerts, runs LLM+RAG analysis per anomaly")
bullet("HEAL     — Executes HealingActions via 21 registered handlers, retries with backoff, sends Slack reports")
bullet("REPORT   — Emits cycle summary to WebSocket event bus")
doc.add_paragraph()
body("Supporting infrastructure:", bold=True)
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Agent loop",        "LangGraph 0.2",                        "State machine orchestration"],
        ["LLM reasoning",     "OpenAI / Bedrock / Groq / Ollama / WatsonX", "Anomaly detection + root-cause analysis"],
        ["RAG retrieval",     "FAISS + LangChain",                    "Runbook context retrieval"],
        ["Embeddings",        "HuggingFace / WatsonX / OpenAI / Ollama / Bedrock", "Semantic search"],
        ["AWS SDK",           "boto3 1.34",                           "Metric collection + healing actions"],
        ["Notifications",     "httpx → Slack Incoming Webhooks",      "Anomaly & healing alerts"],
        ["Orchestrate skills","FastAPI /openapi-orchestrate.json",     "WatsonX Orchestrate integration"],
        ["API server",        "FastAPI + uvicorn",                    "REST + WebSocket + dashboard"],
        ["Config",            "pydantic-settings",                    "Type-safe env-var configuration"],
        ["Structured logs",   "structlog",                            "JSON-formatted operational logs"],
    ],
    col_widths=[1.6, 2.2, 2.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Project Structure")
hr()
code_block(
    "aws-devops-agent/\n"
    "├── main.py                              # Entry point\n"
    "├── requirements.txt                     # Python dependencies\n"
    "├── .env.example                         # Environment variable template\n"
    "├── export_orchestrate_spec.py           # Generate Orchestrate skill file\n"
    "├── openapi-orchestrate.json             # Orchestrate skill descriptor (generated)\n"
    "├── knowledge_base/                      # Runbook Markdown files (RAG source)\n"
    "│   ├── ec2_runbook.md\n"
    "│   ├── rds_runbook.md\n"
    "│   ├── ecs_runbook.md\n"
    "│   ├── lambda_runbook.md\n"
    "│   ├── alb_runbook.md\n"
    "│   ├── general_best_practices.md\n"
    "│   └── learned_fixes.md                 # Auto-updated by POST /api/feedback\n"
    "├── data/\n"
    "│   └── vector_store/                    # FAISS index (auto-created)\n"
    "└── aws_devops_agent/\n"
    "    ├── config.py                        # pydantic-settings configuration\n"
    "    ├── models.py                        # Shared Pydantic data models\n"
    "    ├── server.py                        # FastAPI app + dashboard\n"
    "    ├── agent/\n"
    "    │   └── orchestrator.py              # LangGraph state machine\n"
    "    ├── analyzer/\n"
    "    │   └── engine.py                    # 3-mode anomaly detection + LLM analysis\n"
    "    ├── healing/\n"
    "    │   └── executor.py                  # 21 self-healing action handlers\n"
    "    ├── monitoring/\n"
    "    │   └── collector.py                 # AWS metric collectors\n"
    "    ├── notifications/\n"
    "    │   └── slack.py                     # Slack webhook notifications\n"
    "    └── rag/\n"
    "        └── vector_store.py              # FAISS vector store + retriever"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. INSTALLATION & QUICK START
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Installation & Quick Start")
hr()
heading("4.1  Prerequisites", level=2)
bullet("Python 3.10 or 3.12")
bullet("AWS account with CloudWatch read access + SSM/EC2/RDS/ECS/Lambda/ELBv2 write access")
bullet("One of: OpenAI API key, Groq API key, IBM Cloud API key, AWS Bedrock, or local Ollama")

heading("4.2  Clone & Install", level=2)
code_block(
    "git clone <repo-url>\n"
    "cd aws-devops-agent\n"
    "python -m venv .venv\n"
    "source .venv/bin/activate      # Windows: .venv\\Scripts\\activate\n"
    "pip install -r requirements.txt"
)

heading("4.3  Configure", level=2)
code_block(
    "cp .env.example .env\n"
    "# Fill in: AWS credentials, LLM_PROVIDER + key, ANOMALY_DETECTION_MODE\n"
    "# Optional: SLACK_WEBHOOK_URL for Slack notifications"
)

heading("4.4  Run", level=2)
code_block("python main.py")
body("Open http://localhost:8000 to see the live dashboard.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Configuration Reference")
hr()
body("All configuration is loaded from environment variables or a .env file.")
doc.add_paragraph()
add_table(
    ["Variable", "Default", "Description"],
    [
        ["AWS_ACCESS_KEY_ID",            "(empty)",     "AWS access key (optional if using IAM role)"],
        ["AWS_SECRET_ACCESS_KEY",        "(empty)",     "AWS secret key"],
        ["AWS_DEFAULT_REGION",           "us-east-1",   "AWS region to monitor"],
        ["LLM_PROVIDER",                 "openai",      "openai | bedrock | groq | ollama | watsonx"],
        ["OPENAI_API_KEY",               "(empty)",     "OpenAI API key"],
        ["LLM_MODEL",                    "gpt-4o",      "OpenAI model name"],
        ["BEDROCK_MODEL_ID",             "claude-3-5",  "Bedrock model ID"],
        ["GROQ_API_KEY",                 "(empty)",     "Groq API key"],
        ["GROQ_MODEL",                   "llama-3.3-70b-versatile", "Groq model"],
        ["OLLAMA_BASE_URL",              "http://localhost:11434", "Ollama URL"],
        ["OLLAMA_MODEL",                 "llama3.1",    "Ollama model"],
        ["OLLAMA_EMBED_MODEL",           "nomic-embed-text", "Ollama embedding model"],
        ["EMBED_PROVIDER",               "huggingface", "huggingface | openai | ollama | bedrock | watsonx"],
        ["HUGGINGFACE_API_KEY",          "(empty)",     "HuggingFace read token"],
        ["HF_EMBED_MODEL",               "all-MiniLM-L6-v2", "Sentence-transformer model"],
        ["WATSONX_API_KEY",              "(empty)",     "IBM Cloud API key"],
        ["WATSONX_PROJECT_ID",           "(empty)",     "WatsonX.ai project ID"],
        ["WATSONX_URL",                  "https://us-south.ml.cloud.ibm.com", "WatsonX endpoint"],
        ["WATSONX_LLM_MODEL",            "ibm/granite-3-8b-instruct", "WatsonX LLM model"],
        ["WATSONX_EMBED_MODEL",          "ibm/slate-125m-english-rtrvr", "WatsonX embedding model"],
        ["ANOMALY_DETECTION_MODE",       "hybrid",      "threshold | llm | hybrid"],
        ["SLACK_WEBHOOK_URL",            "(empty)",     "Slack Incoming Webhook URL for alerts"],
        ["AGENT_POLL_INTERVAL_SECONDS",  "60",          "Monitoring cycle interval in seconds"],
        ["AGENT_MAX_HEALING_RETRIES",    "3",           "Max retry attempts per healing action"],
        ["AGENT_DRY_RUN",                "false",       "If true, log actions without executing"],
        ["VECTOR_STORE_PATH",            "./data/vector_store", "Path to FAISS index"],
        ["KNOWLEDGE_BASE_PATH",          "./knowledge_base",    "Runbook directory"],
        ["API_HOST",                     "0.0.0.0",     "FastAPI server bind address"],
        ["API_PORT",                     "8000",        "FastAPI server port"],
        ["MONITOR_EC2_INSTANCE_IDS",     "(empty)",     "Comma-separated EC2 IDs (blank = all)"],
        ["MONITOR_RDS_CLUSTER_IDS",      "(empty)",     "Comma-separated RDS IDs (blank = all)"],
        ["MONITOR_ECS_CLUSTER_NAMES",    "(empty)",     "Comma-separated ECS cluster names"],
        ["MONITOR_LAMBDA_FUNCTION_NAMES","(empty)",     "Comma-separated Lambda names"],
        ["MONITOR_ALB_NAMES",            "(empty)",     "Comma-separated ALB names"],
    ],
    col_widths=[2.3, 1.5, 2.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. CORE MODULES
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Core Modules")
hr()

heading("6.1  config.py — Settings", level=2)
body("Central pydantic-settings singleton. Loaded once via get_settings() with @lru_cache.")
bullet("All settings are read from .env or environment variables")
bullet("New settings added: anomaly_detection_mode, slack_webhook_url, all WatsonX vars")

heading("6.2  models.py — Shared Data Models", level=2)
add_table(
    ["Model", "Key Fields", "Purpose"],
    [
        ["Severity",         "OK | WARNING | CRITICAL | UNKNOWN",                "Anomaly severity level"],
        ["HealingStatus",    "PENDING | IN_PROGRESS | SUCCESS | FAILED | SKIPPED","Action state"],
        ["MetricPoint",      "name, value, unit, timestamp",                      "Single CloudWatch data point"],
        ["ResourceSnapshot", "resource_id, resource_type, region, metrics, raw",  "Point-in-time resource health"],
        ["Anomaly",          "resource_id, severity, title, description, metrics","Detected anomaly"],
        ["HealingAction",    "action_id, anomaly, action_type, parameters, status","Healing action record"],
        ["AgentEvent",       "event_type, payload, timestamp",                    "WebSocket event"],
    ],
    col_widths=[1.8, 2.8, 2.0],
)

heading("6.3  monitoring/collector.py — AWS Metric Collectors", level=2)
body("Polls CloudWatch + service APIs. All collectors are synchronous, called via run_in_executor.")
bullet("collect_ec2()     — CPU, NetworkIn/Out, StatusCheckFailed x3")
bullet("collect_rds()     — CPU, FreeStorageSpace, DatabaseConnections, ReadLatency, WriteLatency, FreeableMemory")
bullet("collect_ecs()     — CPU, MemoryUtilization, RunningCount, DesiredCount")
bullet("collect_lambda()  — Invocations, Errors, Throttles, Duration, ConcurrentExecutions")
bullet("collect_alb()     — RequestCount, 5xx/4xx, TargetResponseTime, HealthyHostCount, UnHealthyHostCount")
bullet("collect_all()     — Calls all five; swallows per-collector exceptions")

heading("6.4  analyzer/engine.py — Anomaly Detection & LLM Analysis", level=2)
body("Three detection functions, mode selected via ANOMALY_DETECTION_MODE in .env:")
bullet("detect_anomalies()        — fast rule-based threshold check against _THRESHOLDS dict")
bullet("detect_anomalies_llm()    — sends ALL metrics to LLM per resource; no fixed thresholds; catches any anomaly the LLM recognises")
bullet("detect_anomalies_hybrid() — threshold first, then LLM for uncovered metrics; deduplicates results")
body("Post-detection (all modes):", bold=True)
bullet("analyze_anomaly() — RAG retrieval (k=6 chunks) + LLM chain -> root_cause, recommended_action, action_parameters, reasoning, confidence")
bullet("_get_llm()        — provider factory: openai | bedrock | groq | ollama | watsonx")

heading("6.5  healing/executor.py — Self-Healing Action Executor", level=2)
body("21 registered action handlers. execute_healing_action() handles DRY_RUN, retries, status updates.")
body("Retry policy: 3 attempts, back-off 30s / 60s / 120s", bold=True)

heading("6.6  rag/vector_store.py — FAISS Vector Store", level=2)
bullet("_get_embeddings() — HuggingFace / OpenAI / Ollama / Bedrock / WatsonX (thread-safe singleton)")
bullet("build_vector_store(force) — load from disk or rebuild from knowledge_base/ docs")
bullet("get_retriever(k=6) — similarity-search retriever")
bullet("Auto-rebuilt by POST /api/knowledge/rebuild and POST /api/feedback")

heading("6.7  agent/orchestrator.py — LangGraph Orchestrator", level=2)
body("Reads ANOMALY_DETECTION_MODE and dispatches to the correct detect function. Wires Slack notifications.")
code_block(
    "monitor  ->  analyze  --(has actions)--> heal --> report\n"
    "                       --(no actions)----------> report"
)
bullet("node_analyze: calls notify_anomaly() per anomaly after detection")
bullet("node_heal:    calls notify_healing() per action after execution, passes analysis summary")

heading("6.8  notifications/slack.py — Slack Webhook Notifications", level=2)
body("Fire-and-forget HTTP POST to SLACK_WEBHOOK_URL using Block Kit attachments.")
bullet("notify_anomaly(anomaly)         — red/yellow card with resource, severity, title, description")
bullet("notify_healing(action, summary) — green/red/yellow card with action, result, reasoning, retry count")
bullet("No-op (debug log only) when SLACK_WEBHOOK_URL is not configured")

heading("6.9  server.py — FastAPI Server", level=2)
body("REST API, WebSocket stream, built-in dashboard, Orchestrate skill descriptor, feedback endpoint.")
bullet("Background tasks at startup: vector store build, event relay, agent loop")
bullet("POST /api/feedback — appends fix to learned_fixes.md, rebuilds FAISS immediately")
bullet("GET /openapi-orchestrate.json — trimmed 3-skill spec for WatsonX Orchestrate import")
bullet("POST /api/heal — on-demand healing trigger (used by Orchestrate skill)")
bullet("POST /api/query — RAG Q&A with LLM-synthesized answer field + sources")

heading("6.10  main.py — Entry Point", level=2)
body("Configures structlog and launches uvicorn. Reads API_HOST and API_PORT from Settings.")
code_block("python main.py")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. ANOMALY DETECTION MODES
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Anomaly Detection Modes")
hr()
body("Set ANOMALY_DETECTION_MODE in .env to choose the detection strategy.")
doc.add_paragraph()
add_table(
    ["Mode", "How it works", "Best for"],
    [
        ["threshold",
         "Compares each metric against a hardcoded (warn, critical, direction) rule in _THRESHOLDS. Fast, deterministic, zero LLM cost.",
         "Production with known SLOs; predictable behaviour"],
        ["llm",
         "Sends ALL metrics for each resource snapshot to the LLM. LLM returns a JSON array of findings with severity and description. No fixed numbers needed.",
         "Unknown workloads; catching uncovered metrics like FreeableMemory, NetworkIn, Duration"],
        ["hybrid (default)",
         "Threshold fast-path first. Then LLM is called on snapshots that have at least one metric not covered by any threshold rule. Results are deduplicated.",
         "Best coverage: known metrics fast + unknown metrics via LLM"],
    ],
    col_widths=[1.1, 3.4, 2.1],
)
doc.add_paragraph()
body("LLM detection prompt instructs the model to detect:", bold=True)
bullet("Sudden spikes or drops in any metric")
bullet("Values that indicate failure states (StatusCheckFailed, UnHealthyHostCount)")
bullet("Values unusual for the resource type (e.g. low FreeableMemory on RDS)")
bullet("Cross-metric combinations that together signal a problem")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. MONITORED RESOURCES & METRICS
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Monitored Resources & Metrics")
hr()
add_table(
    ["Service", "Metrics Collected", "Namespace"],
    [
        ["EC2",    "CPUUtilization, NetworkIn, NetworkOut, StatusCheckFailed, StatusCheckFailed_Instance, StatusCheckFailed_System", "AWS/EC2"],
        ["RDS",    "CPUUtilization, FreeStorageSpace, DatabaseConnections, ReadLatency, WriteLatency, FreeableMemory", "AWS/RDS"],
        ["ECS",    "CPUUtilization, MemoryUtilization, RunningCount, DesiredCount", "AWS/ECS"],
        ["Lambda", "Invocations, Errors, Throttles, Duration, ConcurrentExecutions", "AWS/Lambda"],
        ["ALB",    "RequestCount, HTTPCode_ELB_5XX_Count, HTTPCode_ELB_4XX_Count, HTTPCode_Target_5XX_Count, TargetResponseTime, HealthyHostCount, UnHealthyHostCount, ActiveConnectionCount", "AWS/ApplicationELB"],
    ],
    col_widths=[0.8, 3.8, 2.0],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. SELF-HEALING ACTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("9. Self-Healing Actions (21 total)")
hr()
body("All actions are guarded by DRY_RUN mode. Destructive actions require CRITICAL severity per LLM prompt guidelines.")
doc.add_paragraph()
body("EC2 (6 actions):", bold=True)
add_table(
    ["Action", "What It Does", "Key Parameters"],
    [
        ["restart_ec2_service",  "Restarts systemd service via SSM (first resort for high CPU)",    "instance_id, [service_name]"],
        ["reboot_ec2_instance",  "Reboots EC2 instance via EC2 API (last resort, CRITICAL only)",  "instance_id"],
        ["stop_start_instance",  "Stop + Start EC2 (migrates to new hardware) for status failures", "instance_id"],
        ["scale_out_asg",        "Increases ASG desired capacity (traffic spike)",                  "asg_name, [increment=1]"],
        ["cleanup_disk_ssm",     "Runs /tmp, journal, log cleanup via SSM",                         "instance_id"],
        ["extend_ebs_volume",    "Resizes EBS volume by extra_gb (default +20 GB)",                 "volume_id, [extra_gb=20]"],
    ],
    col_widths=[1.9, 2.8, 1.8],
)
doc.add_paragraph()
body("RDS (3 actions):", bold=True)
add_table(
    ["Action", "What It Does", "Key Parameters"],
    [
        ["reboot_rds_instance", "Reboots RDS DB instance",                                         "db_instance_id"],
        ["modify_rds_storage",  "Increases allocated storage by extra_gb (default +20 GB)",         "db_instance_id, [extra_gb=20]"],
        ["enable_rds_proxy",    "Creates RDS Proxy to pool connections (high DatabaseConnections)", "db_instance_id, role_arn, subnet_ids"],
    ],
    col_widths=[1.9, 2.8, 1.8],
)
doc.add_paragraph()
body("ECS (6 actions):", bold=True)
add_table(
    ["Action", "What It Does", "Key Parameters"],
    [
        ["update_ecs_service",       "Forces a new ECS deployment",                                     "cluster, service"],
        ["scale_ecs_service",        "Increases desired count by increment",                            "cluster, service, [increment=2]"],
        ["rollback_ecs_task_def",    "Redeploys previous task definition revision (bad deploy)",        "cluster, service"],
        ["increase_ecs_task_memory", "Registers new task def with +extra_mb memory, redeploys (OOM)",  "cluster, service, [extra_mb=512]"],
        ["update_ecs_desired_count", "Sets desired count to explicit safe minimum",                     "cluster, service, desired_count"],
        ["toggle_capacity_provider", "Switches FARGATE_SPOT to FARGATE on-demand",                     "cluster, service"],
    ],
    col_widths=[1.9, 2.8, 1.8],
)
doc.add_paragraph()
body("ALB (2 actions):", bold=True)
add_table(
    ["Action", "What It Does", "Key Parameters"],
    [
        ["deregister_unhealthy_targets", "Removes all unhealthy targets from target group", "target_group_arn"],
        ["increase_alb_idle_timeout",    "Increases idle timeout to reduce 504 errors",     "load_balancer_arn, [timeout_seconds=120]"],
    ],
    col_widths=[1.9, 2.8, 1.8],
)
doc.add_paragraph()
body("Lambda (4 actions):", bold=True)
add_table(
    ["Action", "What It Does", "Key Parameters"],
    [
        ["rollback_lambda_version",   "Updates alias to previous version (bad deploy)",          "function_name, [alias=live]"],
        ["update_lambda_timeout",     "Increases timeout by increment_seconds (max 900s)",       "function_name, [increment_seconds=30]"],
        ["update_lambda_memory",      "Increases memory by increment_mb (max 10240 MB)",         "function_name, [increment_mb=256]"],
        ["put_function_concurrency",  "Sets reserved concurrency to prevent throttling",          "function_name, [reserved_concurrency=100]"],
    ],
    col_widths=[1.9, 2.8, 1.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading("10. API Reference")
hr()
add_table(
    ["Method", "Path", "Description"],
    [
        ["GET",  "/",                         "Built-in HTML live dashboard (WatsonX IBM theme)"],
        ["GET",  "/health",                   "Liveness probe — returns {status: ok}"],
        ["GET",  "/api/status",               "Live AWS scan: resources_scanned, anomalies[], dry_run, region"],
        ["POST", "/api/query",                "RAG Q&A -> LLM-synthesized answer + sources[]"],
        ["POST", "/api/heal",                 "On-demand healing action trigger (used by Orchestrate)"],
        ["POST", "/api/feedback",             "Submit operator fix -> learned_fixes.md + FAISS rebuild"],
        ["POST", "/api/knowledge/rebuild",    "Force-rebuild FAISS vector store"],
        ["GET",  "/openapi-orchestrate.json", "WatsonX Orchestrate 3-skill OpenAPI descriptor"],
        ["WS",   "/ws/events",                "WebSocket: streams AgentEvent JSON, ping every 20s"],
    ],
    col_widths=[0.55, 2.0, 4.05],
)
doc.add_paragraph()
body("POST /api/query — request & response:", bold=True)
code_block(
    'Request:  {"question": "How do I fix high RDS CPU?", "k": 6}\n'
    'Response: {"question": "...", "answer": "<LLM answer>", "sources": [...]}'
)
body("POST /api/feedback — request & response:", bold=True)
code_block(
    'Request:  {\n'
    '  "resource_type": "ec2",\n'
    '  "resource_id": "i-0e9efcc6",\n'
    '  "anomaly_title": "EC2 CPUUtilization anomaly",\n'
    '  "agent_action": "restart_ec2_service",\n'
    '  "what_worked": "reboot_ec2_instance",\n'
    '  "root_cause_notes": "nginx memory leak after deploy",\n'
    '  "prevent_recurrence": "Add memory leak check to CI"\n'
    '}\n'
    'Response: {"message": "Feedback recorded and knowledge base rebuilt.", ...}'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. WATSONX ORCHESTRATE INTEGRATION
# ══════════════════════════════════════════════════════════════════════════════
heading("11. WatsonX Orchestrate Integration")
hr()
body(
    "The agent exposes three skills importable into WatsonX Orchestrate via a trimmed "
    "OpenAPI descriptor at GET /openapi-orchestrate.json. Users interact via natural language chat."
)
doc.add_paragraph()
add_table(
    ["Skill (operationId)", "Endpoint", "Triggered when user says..."],
    [
        ["getAwsStatus",        "GET  /api/status", "Check my AWS / Any issues? / What's failing?"],
        ["queryKnowledgeBase",  "POST /api/query",  "How do I fix... / What causes... / Best practice for..."],
        ["triggerHealingAction","POST /api/heal",   "Restart service on... / Scale up... / Reboot RDS..."],
    ],
    col_widths=[2.0, 1.8, 2.8],
)
doc.add_paragraph()
body("Import steps:", bold=True)
bullet("1. Run: python main.py")
bullet("2. Run: C:\\ngrok\\ngrok.exe http 8000   (get https://xxxx.ngrok-free.app URL)")
bullet("3. Run: python export_orchestrate_spec.py https://xxxx.ngrok-free.app")
bullet("4. In WatsonX Orchestrate: Skills catalog -> Add skill -> From file -> upload openapi-orchestrate.json")
bullet("5. Select all 3 skills -> Save -> Publish each -> Add to assistant")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. FEEDBACK & LEARNING LOOP
# ══════════════════════════════════════════════════════════════════════════════
heading("12. Feedback & Learning Loop")
hr()
body(
    "When the agent picks the wrong action or fails to heal, operators submit feedback via "
    "POST /api/feedback. The fix is appended as a structured Markdown entry to "
    "knowledge_base/learned_fixes.md and the FAISS index is immediately rebuilt. "
    "On the next anomaly of the same type, the LLM retrieves the operator's fix as RAG context "
    "and selects the correct action."
)
doc.add_paragraph()
body("Flow:", bold=True)
code_block(
    "Agent fails to heal\n"
    "  -> Operator manually fixes it\n"
    "  -> POST /api/feedback {resource_type, anomaly_title, what_worked, root_cause_notes}\n"
    "  -> Appended to knowledge_base/learned_fixes.md (timestamped Markdown entry)\n"
    "  -> FAISS vector store rebuilt automatically\n"
    "  -> Next cycle: LLM retrieves the learned fix as runbook context\n"
    "  -> Agent selects the correct healing action"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. KNOWLEDGE BASE & RUNBOOKS
# ══════════════════════════════════════════════════════════════════════════════
heading("13. Knowledge Base & Runbooks")
hr()
add_table(
    ["File", "Scenarios Covered"],
    [
        ["ec2_runbook.md",            "High CPU (restart/scale/reboot decision guide), Status Check failed, Disk > 90%"],
        ["rds_runbook.md",            "High DB CPU, FreeStorageSpace < 2 GB, Too many connections (RDS Proxy)"],
        ["ecs_runbook.md",            "CrashLoopBackOff (rollback/memory), Desired count not met, High CPU/Memory"],
        ["lambda_runbook.md",         "Error rate > 5%, Throttles (concurrency), Duration near timeout"],
        ["alb_runbook.md",            "5xx errors, Target response time > 2s, No healthy hosts"],
        ["general_best_practices.md", "Severity classification, healing decision framework, safety guards"],
        ["learned_fixes.md",          "Auto-populated by POST /api/feedback; operator-confirmed fixes"],
    ],
    col_widths=[2.2, 4.4],
)
doc.add_paragraph()
body("Text splitter: chunk_size=800, chunk_overlap=100. Top-6 chunks retrieved per anomaly.", italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. SLACK NOTIFICATION SETUP
# ══════════════════════════════════════════════════════════════════════════════
heading("14. Slack Notification Setup")
hr()
body("Two notification types are sent automatically when SLACK_WEBHOOK_URL is configured:")
doc.add_paragraph()
add_table(
    ["Notification", "When sent", "Content"],
    [
        ["Anomaly alert",  "Immediately when anomaly detected in node_analyze", "Resource ID, type, severity (colour-coded), title, description, timestamp"],
        ["Healing report", "After each healing action completes in node_heal",  "Action type, status (SUCCESS/FAILED/SKIPPED), result message, reasoning, analysis summary, retry count"],
    ],
    col_widths=[1.5, 2.0, 3.1],
)
doc.add_paragraph()
body("To configure Slack:", bold=True)
bullet("1. Go to https://api.slack.com/apps -> Create New App -> From scratch")
bullet("2. Incoming Webhooks -> Activate -> Add New Webhook to Workspace -> choose channel")
bullet("3. Copy webhook URL and add to .env: SLACK_WEBHOOK_URL=https://hooks.slack.com/services/...")
bullet("4. Restart the agent")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. SAFETY & OPERATIONAL GUARDRAILS
# ══════════════════════════════════════════════════════════════════════════════
heading("15. Safety & Operational Guardrails")
hr()
add_table(
    ["Guardrail", "Implementation"],
    [
        ["DRY_RUN mode",            "AGENT_DRY_RUN=true skips all AWS API calls; logs intended action"],
        ["No data deletion",        "No action in ACTION_REGISTRY deletes S3, RDS snapshots, or any data"],
        ["Destructive action gate", "reboot/stop-start require CRITICAL severity (enforced via LLM prompt decision guides)"],
        ["Retry + back-off",        "3 attempts, 30s/60s/120s delays. After final failure: FAILED status logged"],
        ["Event audit trail",       "Every action start and result emitted to WebSocket bus + Slack"],
        ["Unknown action guard",    "LLM recommendation not in ACTION_REGISTRY -> FAILED immediately"],
        ["Queue overflow guard",    "Event bus capped at 1000 entries; oldest dropped on overflow"],
        ["LLM detection fallback",  "If LLM detection throws an exception per resource, it is logged and skipped; other resources continue"],
    ],
    col_widths=[2.0, 4.6],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. EXTENDING THE AGENT
# ══════════════════════════════════════════════════════════════════════════════
heading("16. Extending the Agent")
hr()

heading("16.1  Add a new runbook", level=2)
bullet("Drop a .md file into knowledge_base/ and call POST /api/knowledge/rebuild")

heading("16.2  Add a new healing action", level=2)
bullet("Write _my_action(params: Dict) -> str in healing/executor.py")
bullet("Add 'my_action': _my_action to ACTION_REGISTRY")
bullet("Add the name to the LLM prompt action list in engine.py _ANALYSIS_PROMPT")

heading("16.3  Add a new monitored resource type", level=2)
bullet("Add collect_<type>() in monitoring/collector.py returning List[ResourceSnapshot]")
bullet("Add it to collect_all()")
bullet("Add threshold entries in engine.py _THRESHOLDS (optional — LLM mode covers it automatically)")

heading("16.4  Switch LLM provider", level=2)
bullet("Set LLM_PROVIDER=openai | bedrock | groq | ollama | watsonx in .env and restart")

heading("16.5  Switch detection mode", level=2)
bullet("Set ANOMALY_DETECTION_MODE=threshold | llm | hybrid in .env and restart — no code changes")

heading("16.6  Submit operator feedback", level=2)
bullet("POST /api/feedback with what_worked and root_cause_notes")
bullet("Agent rebuilds FAISS immediately; new fix is retrievable by the next cycle")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17. DEPENDENCIES
# ══════════════════════════════════════════════════════════════════════════════
heading("17. Dependencies")
hr()
add_table(
    ["Package", "Version", "Purpose"],
    [
        ["langchain",               "0.2.16",  "Core LangChain framework"],
        ["langchain-community",     "0.2.16",  "Community integrations (FAISS, etc.)"],
        ["langchain-openai",        "0.1.23",  "OpenAI LLM + embeddings"],
        ["langchain-aws",           "0.1.17",  "AWS Bedrock LLM + embeddings"],
        ["langchain-groq",          "0.1.9",   "Groq LLM"],
        ["langchain-ollama",        "0.1.3",   "Ollama local LLM"],
        ["langchain-huggingface",   "0.0.3",   "HuggingFace embeddings"],
        ["langchain-ibm",           "0.1.12",  "IBM WatsonX LLM + embeddings"],
        ["ibm-watsonx-ai",          "1.1.2",   "IBM WatsonX AI Python SDK"],
        ["langgraph",               "0.2.16",  "LangGraph state machine orchestration"],
        ["openai",                  "1.40.0",  "OpenAI Python SDK"],
        ["groq",                    "0.9.0",   "Groq Python SDK"],
        ["faiss-cpu",               "1.8.0",   "FAISS vector similarity search"],
        ["sentence-transformers",   "3.0.1",   "Local embedding models"],
        ["huggingface-hub",         "0.24.6",  "HuggingFace Hub client"],
        ["boto3",                   "1.34.162","AWS SDK for Python"],
        ["fastapi",                 "0.112.0", "REST + WebSocket API framework"],
        ["uvicorn[standard]",       "0.30.6",  "ASGI server"],
        ["websockets",              "12.0",    "WebSocket protocol"],
        ["pydantic",                "2.8.2",   "Data validation"],
        ["pydantic-settings",       "2.4.0",   "Settings from env vars"],
        ["structlog",               "24.4.0",  "Structured logging"],
        ["httpx",                   "0.27.0",  "HTTP client (Slack webhook calls)"],
        ["tenacity",                "8.5.0",   "Retry / back-off"],
        ["python-dotenv",           "1.0.1",   ".env file loader"],
        ["numpy",                   "1.26.4",  "Numerical ops (FAISS dep)"],
        ["pandas",                  "2.2.2",   "Data utilities"],
    ],
    col_widths=[2.0, 1.0, 3.6],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 18. DATA MODELS REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading("18. Data Models Reference")
hr()
body("All models use Pydantic v2, defined in aws_devops_agent/models.py.")
doc.add_paragraph()

heading("MetricPoint", level=2)
add_table(
    ["Field", "Type", "Default", "Description"],
    [
        ["name",      "str",      "—",        "Metric name (CPUUtilization, Errors, etc.)"],
        ["value",     "float",    "—",        "Numeric metric value"],
        ["unit",      "str",      '""',       "%, Bytes, Count, Seconds, Milliseconds"],
        ["timestamp", "datetime", "utcnow()", "Sampling timestamp"],
    ],
    col_widths=[1.2, 0.9, 1.0, 3.5],
)
doc.add_paragraph()

heading("ResourceSnapshot", level=2)
add_table(
    ["Field", "Type", "Default", "Description"],
    [
        ["resource_id",   "str",               "—",         "EC2 instance ID, function name, etc."],
        ["resource_type", "str",               "—",         "ec2 | rds | ecs | lambda | alb"],
        ["region",        "str",               "—",         "AWS region"],
        ["metrics",       "List[MetricPoint]", "[]",        "Collected metrics"],
        ["tags",          "Dict[str,str]",     "{}",        "AWS resource tags"],
        ["raw",           "Dict[str,Any]",     "{}",        "Raw API fields (instance_type, engine, etc.)"],
        ["collected_at",  "datetime",          "utcnow()",  "Collection timestamp"],
    ],
    col_widths=[1.2, 1.5, 0.7, 3.2],
)
doc.add_paragraph()

heading("Anomaly", level=2)
add_table(
    ["Field", "Type", "Default", "Description"],
    [
        ["resource_id",   "str",               "—",        "Resource that triggered the anomaly"],
        ["resource_type", "str",               "—",        "ec2 | rds | ecs | lambda | alb"],
        ["severity",      "Severity",          "—",        "WARNING or CRITICAL"],
        ["title",         "str",               "—",        "Short description"],
        ["description",   "str",               "—",        "Metric value detail or LLM explanation"],
        ["metrics",       "List[MetricPoint]", "[]",       "Triggering metrics"],
        ["detected_at",   "datetime",          "utcnow()", "Detection timestamp"],
        ["context",       "Dict[str,Any]",     "{}",       "Raw resource context"],
    ],
    col_widths=[1.2, 1.5, 0.7, 3.2],
)
doc.add_paragraph()

heading("HealingAction", level=2)
add_table(
    ["Field", "Type", "Default", "Description"],
    [
        ["action_id",   "str",               "—",       "UUID"],
        ["anomaly",     "Anomaly",           "—",       "Triggering anomaly"],
        ["action_type", "str",               "—",       "One of 21 registered action names"],
        ["parameters",  "Dict[str,Any]",     "{}",      "Action-specific parameters"],
        ["reasoning",   "str",               '""',      "LLM or operator justification"],
        ["status",      "HealingStatus",     "PENDING", "Execution state"],
        ["result",      "Optional[str]",     "None",    "Outcome message from handler"],
        ["executed_at", "Optional[datetime]","None",    "Execution start timestamp"],
        ["retries",     "int",               "0",       "Retry count"],
    ],
    col_widths=[1.2, 1.5, 0.9, 3.0],
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
hr()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    f"AWS DevOps RAG Agent — IBM WatsonX Edition v2  |  {datetime.date.today().strftime('%Y')}"
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)
fr.font.italic = True

output_path = "AWS_DevOps_RAG_Agent_WatsonX_Documentation.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
